import streamlit as st
import pandas as pd
from docx import Document
import zipfile
import io
from datetime import datetime

# Configurar página y tema
st.set_page_config(
    page_title="Automatización de Documentos - Atenea",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# CSS personalizado para el diseño premium
st.markdown("""
<style>
    /* Fondo oscuro */
    body {
        background-color: #0f172a;
    }
    
    /* Hero */
    .hero {
        text-align: center;
        margin-bottom: 40px;
        padding: 40px 0;
    }
    
    .hero h1 {
        font-size: 48px;
        font-weight: 700;
        margin-bottom: 12px;
        background: linear-gradient(135deg, #a855f7 0%, #9333ea 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    /* Container */
    .upload-section {
        background: rgba(15, 23, 42, 0.5);
        border-radius: 20px;
        padding: 50px 40px;
        backdrop-filter: blur(10px);
    }
    
    .upload-grid {
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 24px;
        margin-bottom: 40px;
    }
    
    .upload-box {
        background: rgba(15, 23, 42, 0.8);
        border: 2px dashed rgba(168, 85, 247, 0.3);
        border-radius: 16px;
        padding: 40px;
        text-align: center;
        transition: all 0.3s;
    }
    
    .upload-box:hover {
        border-color: rgba(168, 85, 247, 0.6);
        background: rgba(15, 23, 42, 0.9);
        transform: translateY(-4px);
        box-shadow: 0 20px 40px rgba(168, 85, 247, 0.1);
    }
</style>
""", unsafe_allow_html=True)

# Hero Section
st.markdown("""
<div class="hero">
    <h1>Automatización de Documentos</h1>
    <p style="font-size: 18px; color: #a0aec0; margin-top: 12px;">
        Carga tu Excel y plantilla, genera documentos automáticamente
    </p>
</div>
""", unsafe_allow_html=True)

def replace_text_in_paragraph(paragraph, key, value):
    """Reemplaza placeholders en párrafos manteniendo formato"""
    placeholder = "{{" + key + "}}"
    value_str = str(value) if value is not None else ""
    
    full_text = "".join(run.text for run in paragraph.runs)
    
    if placeholder in full_text:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value_str)
                return

def generar_documentos(df, word_file):
    """Genera documentos Word a partir de Excel"""
    try:
        zip_buffer = io.BytesIO()
        documentos_generados = 0
        errores = []
        
        with zipfile.ZipFile(zip_buffer, "w") as zipf:
            for idx, row in df.iterrows():
                try:
                    word_file.seek(0)
                    doc = Document(word_file)
                    
                    for paragraph in doc.paragraphs:
                        for key in row.index:
                            replace_text_in_paragraph(paragraph, key, row[key])
                    
                    for table in doc.tables:
                        for row_table in table.rows:
                            for cell in row_table.cells:
                                for paragraph in cell.paragraphs:
                                    for key in row.index:
                                        replace_text_in_paragraph(paragraph, key, row[key])
                    
                    doc_bytes = io.BytesIO()
                    doc.save(doc_bytes)
                    zipf.writestr(f"Documento_{idx + 1}.docx", doc_bytes.getvalue())
                    documentos_generados += 1
                    
                except Exception as e:
                    errores.append(f"Fila {idx + 1}: {str(e)}")
                    
        return zip_buffer, documentos_generados, errores
        
    except Exception as e:
        raise Exception(f"Error procesando documentos: {str(e)}")

# Sección de carga
col1, col2 = st.columns(2)

with col1:
    excel_file = st.file_uploader("📊 Carga tu EXCEL", type="xlsx")
with col2:
    word_file = st.file_uploader("📄 Carga tu PLANTILLA", type="docx")

if excel_file and word_file:
    st.success("✅ Archivos cargados correctamente")
    
    with st.expander("👁️ Ver datos del Excel"):
        try:
            df_preview = pd.read_excel(excel_file)
            st.dataframe(df_preview, use_container_width=True)
            st.info(f"Total de filas: {len(df_preview)}")
        except Exception as e:
            st.error(f"Error al leer Excel: {e}")
    
    if st.button("⚡ Generar Documentos", use_container_width=True, type="primary"):
        try:
            excel_file.seek(0)
            df = pd.read_excel(excel_file)
            
            if df.empty:
                st.error("El Excel está vacío")
            elif len(df) > 1000:
                st.warning(f"Tienes {len(df)} filas. Esto puede tardar...")
            
            with st.spinner(f"Procesando {len(df)} documentos..."):
                zip_buffer, generados, errores = generar_documentos(df, word_file)
                
                if generados > 0:
                    st.success(f"✅ Se generaron {generados} documentos correctamente")
                    
                    if errores:
                        with st.expander(f"⚠️ Errores en {len(errores)} documentos"):
                            for error in errores:
                                st.warning(error)
                    
                    zip_buffer.seek(0)
                    st.download_button(
                        label=f"📥 Descargar ZIP ({generados} documentos)",
                        data=zip_buffer.getvalue(),
                        file_name=f"Documentos_Generados_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                        mime="application/zip",
                        use_container_width=True,
                        type="primary"
                    )
                else:
                    st.error("No se pudieron generar documentos. Revisa los errores abajo.")
                    with st.expander("Ver errores"):
                        for error in errores:
                            st.error(error)
                            
        except Exception as e:
            st.error(f"Error: {str(e)}")
            st.info("Verifica que: 1) El Excel tenga datos, 2) La plantilla tenga placeholders {{columna}}")
            
else:
    st.warning("⚠️ Por favor carga ambos archivos para continuar")
    st.info("📋 Pasos: 1) Carga un Excel con los datos, 2) Carga una plantilla Word con placeholders {{columna}}, 3) Haz clic en Generar")
